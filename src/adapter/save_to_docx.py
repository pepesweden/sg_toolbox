#Funktion för att formatera och  spara summering i word

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Laddar json bibliotek
import json

#Definera mappar för static content


# Definierar punktlista definerar punklista
def add_bullet_list(items, doc):
    for item in items:
        paragraph = doc.add_paragraph(style='List Bullet')
        run = paragraph.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)

def save_summary_to_docx(response_text, filename):
    """
    Parsar AI-svar UPPIFRÅN OCH NER.
    Identifierar varje element och formaterar det.
    """
    doc = Document()
    
    lines = response_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Test 1: Är det en titel?
        if line.startswith('# ') or line.startswith('**# '):
            print("✅ Huvudrubkrik identidierad")
            add_title(doc, line)
            i += 1
        
        # Test 2: Är det JSON-start?
        elif '[Start Json]' in line:
            i = add_json_table(doc, lines, i)  # Returnerar nytt index
        
        # Test 3: Är det underrubrik?
        elif line.startswith('## ') or line.startswith('**## ') or line.startswith('**'):
            add_heading(doc, line)
            i += 1
        
        # Test 4: Normal text
        else:
            add_paragraph(doc, line)
            i += 1
    
    doc.save(filename)


# Formatterar och Sparar sammanfattning i en word fil
def save_summary_to_docx_OLD(summary_text, candidate_name, filepath):
    doc = Document()

    # Lägg till logotyp i sidhuvud
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    run = header_paragraph.add_run()
    run.add_picture("ui/static/logo/standard.png", width=Inches(1.5))
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    #  Extrahera rubriken (första raden)
    lines = summary_text.strip().split("\n")
    title_line = lines[0].strip()
    remaining_lines = lines[1:]

    # Lägg till rubriken
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title_line)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = "Arial"
    title_run.font.color.rgb = RGBColor(106, 13, 173)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Första stycket → punktlista
    paragraphs = "\n".join(remaining_lines).strip().split("\n\n")
    bullet_text = paragraphs[0]
    bullet_points = [line.strip("-• ") for line in bullet_text.split("\n") if line.strip()]
    add_bullet_list(bullet_points, doc)

    # Formatera resterande innehåll
    for para in "\n".join(paragraphs[1:]).split("\n"):
        if not para.strip():
            continue

        paragraph = doc.add_paragraph()
        
        # Rubrik 2 (mindre rubriker)
        if para.strip().startswith("##"): #and para.strip().endswith("**"):
            clean_text = para.strip().replace("##", "")
            run = paragraph.add_run(clean_text)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = "Arial"
            run.font.color.rgb = RGBColor(106, 13, 173)
        # Brödtext
        else:
            run = paragraph.add_run(para.strip())
            run.font.size = Pt(10)
            run.font.name = "Arial"
            run.font.color.rgb = RGBColor(0, 0, 0)

    #  Spara fil
    # filename = f"data/output/sammanfattning_{candidate_name.lower().replace(' ', '_')}.docx"
    doc.save(filepath)

#####################################
###   Creat and format table.     ###
#####################################

# Function to set cell backgroudnd colour
def set_cell_background(cell, color):
    """Sätter bakgrundsfärg på cell (hex-färg)"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading)

#Create and format table from JSON
def create_table_from_json(doc, table_data):
    """
    Skapar en formaterad tabell från JSON-data.
    
    Args:
        doc: python-docx Document object
        table_data: Dict med "tabell"-key
    """
    # Steg 1: Skapa tom tabell (5 rader, 4 kolumner)
    table = doc.add_table(rows=5, cols=4)
    # Add grid to table, this is changing an element not paragraph formatting
    table.style = 'Table Grid'

    # TODO: Fyll i data (nästa steg)
    row = 0
    merged_row = 3 
    for row_object in table_data["tabell"]:
        print(f"--- Rad {row} börjar ---")
        col = 0 # <--- col måste återställar inför varje loop över items i listan "tabell"
        for item in row_object["cells"]:
            if row_object["merged"] == True:
                print(f"Sätter cell ({row}, {col})")
                #table.cell(row, col).text = item["label"]
                cell = table.cell(row, col)
                set_cell_background(cell, "EDEDED")
                cell.text = item["label"]
                para = cell.paragraphs[0]
                run = cell.paragraphs[0].runs[0]
                #Formattering
                run.font.name = "Arial"
                run.font.bold = True
                run.font.size = Pt(10)
            
                para.paragraph_format.space_before = Pt(3)  
                para.paragraph_format.space_after = Pt(3)


                col += 1
                # print(f"Sätter cell ({row}, {col})") # <- To se the cellcretion is done according to plan
                #table.cell(row, col).text = item["value"]
                cell_value = table.cell(row, col)
                cell_value.text = item["value"]
                value_para = cell_value.paragraphs[0]
                value_run = cell_value.paragraphs[0].runs[0]
                #Formattering
                value_run.font.name = "Arial"
                value_run.font.size = Pt(10)
            
                value_para.paragraph_format.space_before = Pt(3)  
                value_para.paragraph_format.space_after = Pt(3)
                col += 1
                table.cell(merged_row, 1).merge(table.cell(merged_row, 3))
                merged_row += 1
            else:
                print(f"Sätter cell ({row}, {col})")
                #table.cell(row, col).text = item["label"]
                cell = table.cell(row, col)
                set_cell_background(cell, "EDEDED")
                cell.text = item["label"]
                para = cell.paragraphs[0]
                run = cell.paragraphs[0].runs[0] # <--- Picks pout the first paragraph in the cell and the deafault run index 0
                #Formattering
                run.font.name = "Arial"
                run.font.bold = True
                run.font.size = Pt(10)
            
                para.paragraph_format.space_before = Pt(3)  
                para.paragraph_format.space_after = Pt(3)
                
                col += 1
                #table.cell(row, col).text = item["value"]
                cell_value = table.cell(row, col)
                cell_value.text = item["value"]
                value_para = cell_value.paragraphs[0]
                value_run = cell_value.paragraphs[0].runs[0]
                #Formattering
                value_run.font.name = "Arial"
                value_run.font.size = Pt(10)
            
                value_para.paragraph_format.space_before = Pt(3)  
                value_para.paragraph_format.space_after = Pt(3)
                col += 1
        row += 1
    return table


#####################################
### Text Formatting functions     ###
#####################################

def add_title(doc, line):
    text = line.replace("# ", "").replace("**", "")
    heading = doc.add_heading(text, level=0)
    heading.runs[0].font.size = Pt(20)
    heading.runs[0].font.bold = True
    heading.runs[0].font.name = "Arial"
    heading.runs[0].font.color.rgb = RGBColor(69, 40, 111) 
    print("✅ Huvudrubrik formatterad")
   

def add_heading(doc, line):
    text = line.replace('## ', '').replace('** ', '')
    para = doc.add_paragraph(text)
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.bold = True
    para.runs[0].font.name = "Arial"
    para.runs[0].font.color.rgb = RGBColor(106, 43, 113)

def add_json_table(doc, lines, start_index):
    """
    Extraherar JSON från lines, skapar tabell.
    Returnerar index EFTER [SLUT Json]
    """
    # Samla JSON-rader
    json_lines = [] 
    i = start_index + 1
    
    while i < len(lines):
        if '[SLUT Json]' in lines[i]:
            break
        json_lines.append(lines[i])
        i += 1
    
    # Parsa JSON
    json_str = '\n'.join(json_lines)
    json_str = json_str.replace('{{', '{').replace('}}', '}')
    table_data = json.loads(json_str)
    
    # Skapa tabell
    create_table_from_json(doc, table_data)
    
    return i + 1  # Returnera index efter [SLUT Json]

def add_paragraph(doc, line):
    if line:  # Skippa tomma rader
        para = doc.add_paragraph(line)
        para.runs[0].font.size = Pt(9.5)
        para.runs[0].font.name = "Arial"

"""
    # Sparar kp i en word fil
def save_kp_to_docx(summary_text, candidate_name):
    doc = Document()

    # Lägg till logotyp i sidhuvud
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    run = header_paragraph.add_run()
    run.add_picture("static/logo/standard.png", width=Inches(1.5))
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    #  Extrahera rubriken (första raden)
    lines = summary_text.strip().split("\n")
    title_line = lines[0].strip()
    remaining_lines = lines[1:]

    # Lägg till rubriken
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title_line)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = "Arial"
    title_run.font.color.rgb = RGBColor(106, 13, 173)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Första stycket → punktlista
    paragraphs = "\n".join(remaining_lines).strip().split("\n\n")
    bullet_text = paragraphs[0]
    bullet_points = [line.strip("-• ") for line in bullet_text.split("\n") if line.strip()]
    add_bullet_list(bullet_points, doc)

    # Formatera resterande innehåll
    for para in "\n".join(paragraphs[1:]).split("\n"):
        if not para.strip():
            continue

        paragraph = doc.add_paragraph()
        
        # Rubrik 2 (mindre rubriker)
        if para.strip().startswith("**") and para.strip().endswith("**"):
            clean_text = para.strip().replace("**", "")
            run = paragraph.add_run(clean_text)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = "Arial"
            run.font.color.rgb = RGBColor(106, 13, 173)
        # Brödtext
        else:
            run = paragraph.add_run(para.strip())
            run.font.size = Pt(10)
            run.font.name = "Arial"
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Spara fil
    filename = f"output/kp_{candidate_name.lower().replace(' ', '_')}.docx"
    doc.save(filename)
"""

"""
def save_refsum_to_docx(summary_text, candidate_name):
    doc = Document()

    #  Lägg till logotyp i sidhuvud
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    run = header_paragraph.add_run()
    run.add_picture("static/logo/standard.png", width=Inches(1.5))
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    #  Extrahera rubriken (första raden)
    lines = summary_text.strip().split("\n")
    title_line = lines[0].strip()
    remaining_lines = lines[1:]

    # Lägg till rubriken
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title_line)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = "Arial"
    title_run.font.color.rgb = RGBColor(106, 13, 173)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Första stycket → punktlista
    paragraphs = "\n".join(remaining_lines).strip().split("\n\n")
    bullet_text = paragraphs[0]
    bullet_points = [line.strip("-• ") for line in bullet_text.split("\n") if line.strip()]
    add_bullet_list(bullet_points, doc)

    # Formatera resterande innehåll
    for para in "\n".join(paragraphs[1:]).split("\n"):
        if not para.strip():
            continue

        paragraph = doc.add_paragraph()
        
        # Rubrik 2 (mindre rubriker)
        if para.strip().startswith("**") and para.strip().endswith("**"):
            clean_text = para.strip().replace("**", "")
            run = paragraph.add_run(clean_text)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = "Arial"
            run.font.color.rgb = RGBColor(106, 13, 173)
        # Brödtext
        else:
            run = paragraph.add_run(para.strip())
            run.font.size = Pt(10)
            run.font.name = "Arial"
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Spara fil
    filename = f"output/refsum_{candidate_name.lower().replace(' ', '_')}.docx"
    doc.save(filename)
"""

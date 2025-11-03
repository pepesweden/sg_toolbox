#This Module...
#Loads the docx as input
#Extraxts the text


#Laddar docx modulen
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Union
from pathlib import Path
from werkzeug.datastructures import FileStorage

#Laddar moduler för PDF läsning, OCR och textigenkänning
from PyPDF2 import PdfReader
import pytesseract
from pdf2image import convert_from_path

# Läser in en(1) Word-fil och extraherar texten
def read_docx_text(file_path):
    doc = Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])

# Opens file and extracts raw-text content 
def read_md_text(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()
    
###ny funktion som extraherar text från en eller flera docx-filer
def extract_texts_from_docx(files: List[Union[str, Path, FileStorage]]) -> List[str]:
    results = []

    for file in files:
        try:
            filename = file.filename if isinstance(file, FileStorage) else str(file)
            filename = filename.lower()

            if filename.endswith(".docx"):
                doc = Document(file)
                text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
                if text:
                    results.append(text.strip())

        except Exception as e:
            print(f"[FEL] Kunde inte läsa {filename}: {e}")

    return results


def read_pdf_text(file_path):    # NY
    """Läser .pdf (med OCR fallback)"""
    # Försök extrahera text
    reader = PdfReader(file_path)
    page = reader.pages[0]
  
    #print(page.extract_text())
    if len(page.extract_text()) < 50:
        text = ""
        pages_path = file_path
        pages = convert_from_path(pages_path, 300)

        for i, page_image in enumerate(pages, start=1):
            print(f"🔍 OCR bearbetar sida {i}/{len(pages)}...", end='\r')
            text += pytesseract.image_to_string(page_image, lang="eng")
        print("📄 Extraherade text från bilder")
        return text 
    else:
        text = ""
        # Loop över ALLA sidor och samla text
        for page in reader.pages:  # Loopa över alla sidor
            text += page.extract_text()
        print("📄 Extraherade ren text")
        return text